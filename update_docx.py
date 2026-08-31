# -*- coding: utf-8 -*-
import docx

# Open the existing English docx
doc_path = r'C:\Users\xiaomi\Desktop\igem\BrewXOS_Small_Smart_Fermentation_System_Hardware_Recommendations.docx'
doc = docx.Document(doc_path)

# Update Section II (table 0): Base controller row
# In our doc, table 0 has 7 rows: header + 6 data rows
# Row 3 (index 3) is "Base controller"
t0 = doc.tables[0]
# Confirm by content
print("Before update, table 0 row 3:")
for ci, cell in enumerate(t0.rows[3].cells):
    print(f"  C{ci}: {cell.text[:100]}")

# Update base controller cell (column 1)
new_base = ('DFRduino Mega2560 (ATmega2560, 16 MHz, 256 KB Flash, '
            '54 digital I/O, 16 analog in, 4 UARTs) + I2C LCD / buttons + '
            'SD card module / serial data export; optional ESP8266 for Wi-Fi upload. '
            'Raspberry Pi option for graphical UI on V3.')
# Need to clear existing content first
t0.rows[3].cells[1].text = ''
t0.rows[3].cells[1].paragraphs[0].add_run(new_base)

# Update Section VII (table 3): Control & display row
# In our doc, table 3 (cost table) is the 4th table
# Row 9 (index 9) is "Control & display"
t_cost = doc.tables[3]
print("Before update, cost table row 9:")
for ci, cell in enumerate(t_cost.rows[9].cells):
    print(f"  C{ci}: {cell.text[:80]}")

new_control = ('DFRduino Mega2560 (~RMB 150-250), I2C LCD (20x4, ~RMB 60-120), '
               'SD card module (~RMB 20-40), buttons, 5V/3.3V PSU; add ESP8266 '
               '(~RMB 30-50) for Wi-Fi upload on V2. Raspberry Pi raises cost significantly.')
t_cost.rows[9].cells[3].text = ''
t_cost.rows[9].cells[3].paragraphs[0].add_run(new_control)

# Also update the header row 9 column 0 to make it more descriptive
# (Keep it as "Control & display" - no change)

# Save
doc.save(doc_path)
print(f'\nUpdated docx: {doc_path}')

# Verify
doc2 = docx.Document(doc_path)
print("\n=== Verification ===")
print("Section II - Base controller:")
print(f"  {doc2.tables[0].rows[3].cells[1].text[:150]}")
print("\nSection VII - Control & display notes:")
print(f"  {doc2.tables[3].rows[9].cells[3].text[:150]}")
